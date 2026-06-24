#!/usr/bin/env python3
"""
Locus Vita — Gerador de SPEC (.docx)
Uso: python3 generate_spec.py --context context.json --output SPEC_Feature_v1.0.docx
"""

import argparse
import json
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


# ── Paleta Locus Vita ──────────────────────────────────────────────────────────
AZUL_PROFUNDO  = RGBColor(0x1A, 0x3A, 0x5C)   # #1A3A5C
VERDE_MENTA    = RGBColor(0x78, 0xC2, 0xAD)   # #78C2AD
CINZA_CLARO    = RGBColor(0xF2, 0xF0, 0xEB)   # #F2F0EB — background
TEXTO_ESCURO   = RGBColor(0x1F, 0x29, 0x37)   # quase-preto
CINZA_MEDIO    = RGBColor(0xE5, 0xE7, 0xEB)   # bordas de tabela


def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de uma célula via XML direto."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def set_cell_borders(table, border_color="E5E7EB"):
    """Aplica bordas finas em todas as células da tabela."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), border_color)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def style_header_row(row, bg="#1A3A5C"):
    """Estiliza a linha de cabeçalho de uma tabela."""
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)


def add_table_with_headers(doc, headers: list[str], col_widths: list[float] | None = None) -> object:
    """Cria uma tabela com cabeçalho estilizado."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_cell_borders(table)

    if col_widths:
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = Cm(col_widths[i])

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    return table


def add_data_row(table, values: list[str], bg: str | None = None):
    """Adiciona uma linha de dados à tabela."""
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = TEXTO_ESCURO
        if bg:
            set_cell_bg(cell, bg)
    return row


def add_heading(doc, text: str, level: int):
    """Adiciona cabeçalho com estilo Locus Vita."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if level == 1:
            run.font.color.rgb = AZUL_PROFUNDO
            run.font.size = Pt(14)
            run.bold = True
        elif level == 2:
            run.font.color.rgb = VERDE_MENTA
            run.font.size = Pt(11)
            run.bold = True
        else:
            run.font.color.rgb = TEXTO_ESCURO
            run.font.size = Pt(10)
            run.bold = True
    return p


def add_paragraph(doc, text: str, size: int = 10, italic: bool = False, color: RGBColor = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color or TEXTO_ESCURO
    return p


def build_spec(ctx: dict, output_path: str):
    doc = Document()

    # ── Configuração da página ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin  = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin   = Cm(2)
    section.bottom_margin = Cm(2)

    # ── Estilos base ────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = TEXTO_ESCURO

    meta = ctx.get("meta", {})

    # ── Cabeçalho ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Especificação de Funcionalidade")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = AZUL_PROFUNDO

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(
        f"Produto: {meta.get('produto','Locus Vita')}  ·  "
        f"Feature: {meta.get('feature','—')}  ·  "
        f"v{meta.get('versao','1.0')}  ·  {meta.get('data', datetime.today().strftime('%m/%Y'))}"
    )
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run2.italic = True

    doc.add_paragraph()

    # Tabela de metadados
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"
    set_cell_borders(meta_table)

    meta_rows = [
        ("Produto",          meta.get("produto", "Locus Vita")),
        ("Módulo / Feature", meta.get("feature", "—")),
        ("Autor",            meta.get("autor", "Fábio")),
        ("Data",             meta.get("data", datetime.today().strftime("%d/%m/%Y"))),
        ("Versão",           meta.get("versao", "1.0")),
        ("Status",           meta.get("status", "Planejado")),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_bg(row.cells[0], "1A3A5C")
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)

    doc.add_paragraph()

    # ── 1. VISÃO GERAL ──────────────────────────────────────────────────────
    add_heading(doc, "1. VISÃO GERAL", 1)
    visao = ctx.get("visao_geral", {})
    resumo = visao.get("resumo", "")
    if resumo:
        add_paragraph(doc, resumo)

    kpis = visao.get("kpis", [])
    if kpis:
        doc.add_paragraph()
        add_paragraph(doc, "Métricas-chave de sucesso:", size=10)
        t = add_table_with_headers(doc, ["ID", "Métrica", "Meta"], [1.5, 8, 4])
        for kpi in kpis:
            add_data_row(t, [kpi.get("id", ""), kpi.get("metrica", ""), kpi.get("meta", "")])

    # 1.2 Impacto em Saúde Esperado
    impacto_saude = visao.get("impacto_saude", {})
    comportamento_alvo = impacto_saude.get("comportamento_alvo", "")
    kpis_saude = impacto_saude.get("kpis_saude", [])
    if comportamento_alvo or kpis_saude:
        add_heading(doc, "1.2 Impacto em Saúde Esperado", 2)
        if comportamento_alvo:
            add_paragraph(doc, comportamento_alvo)
        if kpis_saude:
            doc.add_paragraph()
            t = add_table_with_headers(doc, ["Indicador de Saúde", "Baseline Estimado", "Meta"], [5.5, 4, 4])
            for ks in kpis_saude:
                add_data_row(t, [ks.get("indicador",""), ks.get("baseline","—"), ks.get("meta","")])
        doc.add_paragraph()

    # 1.3 Escopo (renomeado de 1.1)
    add_heading(doc, "1.3 Escopo", 2)
    escopo = ctx.get("escopo", {})
    itens_escopo = escopo.get("incluido", []) + escopo.get("excluido", [])
    if itens_escopo:
        t = add_table_with_headers(doc, ["", "Funcionalidade", "Status", "Observação"], [0.6, 7, 3, 3.5])
        for item in itens_escopo:
            add_data_row(t, [
                item.get("icone", ""),
                item.get("funcionalidade", ""),
                item.get("status", ""),
                item.get("obs", ""),
            ])

    doc.add_paragraph()

    # ── 2. PROBLEMA ─────────────────────────────────────────────────────────
    add_heading(doc, "2. PROBLEMA", 1)
    problema = ctx.get("problema", "")
    if problema:
        add_paragraph(doc, problema)

    # 2.1 Contexto Comportamental (COM-B)
    com_b = ctx.get("contexto_comportamental", {})
    barreira = com_b.get("barreira_principal", "")
    analise = com_b.get("analise", "")
    estrategia = com_b.get("estrategia_solucao", "")
    if barreira or analise:
        doc.add_paragraph()
        add_heading(doc, "2.1 Contexto Comportamental (COM-B)", 2)
        if analise:
            add_paragraph(doc, analise)
        if barreira or estrategia:
            doc.add_paragraph()
            t = add_table_with_headers(doc, ["Dimensão COM-B", "Avaliação", "Estratégia da Feature"], [3.5, 5, 6])
            perna = com_b.get("pernas", [])
            if perna:
                for p in perna:
                    add_data_row(t, [p.get("dimensao",""), p.get("avaliacao",""), p.get("estrategia","")])
            else:
                add_data_row(t, [barreira, analise, estrategia])

    doc.add_paragraph()

    # ── 3. USER STORIES ─────────────────────────────────────────────────────
    add_heading(doc, "3. USER STORIES", 1)
    stories = ctx.get("user_stories", [])
    if stories:
        t = add_table_with_headers(doc, ["ID", "User Story", "Critério de Aceite", "Prioridade"], [1.2, 6.5, 5, 1.8])
        for s in stories:
            add_data_row(t, [
                s.get("id", ""),
                s.get("story", ""),
                s.get("criterio", ""),
                s.get("prioridade", ""),
            ])
    doc.add_paragraph()

    # ── 4. FLUXO PRINCIPAL ──────────────────────────────────────────────────
    add_heading(doc, "4. FLUXO PRINCIPAL", 1)
    fluxo = ctx.get("fluxo_principal", {})
    passos = fluxo.get("passos", [])
    for i, passo in enumerate(passos, 1):
        add_paragraph(doc, f"{i}. {passo}")

    alts = fluxo.get("fluxos_alternativos", [])
    if alts:
        doc.add_paragraph()
        add_paragraph(doc, "Fluxos alternativos:", size=10)
        for alt in alts:
            add_paragraph(doc, f"• {alt.get('alternativa','')}: {alt.get('tratamento','')}")

    bordas = fluxo.get("casos_de_borda", [])
    if bordas:
        doc.add_paragraph()
        add_paragraph(doc, "Casos de borda:", size=10)
        t = add_table_with_headers(doc, ["Situação", "Comportamento Esperado", "Estado / Mensagem"], [4, 6, 4])
        for b in bordas:
            add_data_row(t, [b.get("situacao",""), b.get("comportamento",""), b.get("estado","")])

    doc.add_paragraph()

    # ── 5. REGRAS DE NEGÓCIO ────────────────────────────────────────────────
    add_heading(doc, "5. REGRAS DE NEGÓCIO", 1)
    rns = ctx.get("regras_negocio", [])
    if rns:
        t = add_table_with_headers(doc, ["ID", "Regra", "Lógica / Fórmula", "Exemplo"], [1.2, 4, 5, 4])
        for rn in rns:
            add_data_row(t, [rn.get("id",""), rn.get("regra",""), rn.get("logica",""), rn.get("exemplo","")])
    else:
        add_paragraph(doc, "Sem regras determinísticas específicas além dos padrões do projeto.", italic=True, color=RGBColor(0x9C, 0xA3, 0xAF))
    doc.add_paragraph()

    # ── 6. REQUISITOS NÃO-FUNCIONAIS ────────────────────────────────────────
    add_heading(doc, "6. REQUISITOS NÃO-FUNCIONAIS", 1)
    rnfs = ctx.get("requisitos_nao_funcionais", [])
    if rnfs:
        t = add_table_with_headers(doc, ["Categoria", "Requisito", "Critério de Aceite"], [3, 6.5, 5])
        for rnf in rnfs:
            add_data_row(t, [rnf.get("categoria",""), rnf.get("requisito",""), rnf.get("criterio","")])
    doc.add_paragraph()

    # ── 7. ARQUITETURA ──────────────────────────────────────────────────────
    add_heading(doc, "7. ARQUITETURA", 1)
    arq = ctx.get("arquitetura", {})

    add_heading(doc, "7.1 Componentes Principais", 2)
    comps = arq.get("componentes", [])
    if comps:
        t = add_table_with_headers(doc, ["Arquivo", "Responsabilidade", "Ação"], [6, 6.5, 2])
        for comp in comps:
            acao = comp.get("acao", "Modificar")
            bg = "E8F5E9" if acao == "Criar" else None  # verde claro para novos arquivos
            row = add_data_row(t, [comp.get("arquivo",""), comp.get("responsabilidade",""), acao], bg=bg)

    add_heading(doc, "7.2 Modelo de Dados", 2)
    dados = arq.get("modelo_dados", [])
    if dados:
        t = add_table_with_headers(doc, ["Entidade / Coluna", "Tipo", "Descrição"], [4, 2, 8])
        for d in dados:
            add_data_row(t, [d.get("entidade_coluna",""), d.get("tipo",""), d.get("descricao","")])

    add_heading(doc, "7.3 Migration", 2)
    requer = arq.get("requer_migration", False)
    nome_mig = arq.get("nome_migration", "")
    if requer:
        add_paragraph(doc, f"✅ Requer migration: {nome_mig if nome_mig else '[nome_da_migration]'}")
    else:
        add_paragraph(doc, "❌ Não requer migration de banco de dados.")

    # 7.4 Padrão Visual — somente se a feature introduz novos componentes de UI
    padrao_visual = arq.get("padrao_visual", [])
    if padrao_visual:
        doc.add_paragraph()
        add_heading(doc, "7.4 Padrão Visual", 2)
        add_paragraph(
            doc,
            "Tokens e classes para os novos componentes de UI desta feature. "
            "Regras globais: nunca hardcode hex/cores — sempre tokens semânticos de src/index.css.",
            size=9,
            italic=True,
            color=RGBColor(0x9C, 0xA3, 0xAF)
        )
        doc.add_paragraph()
        t = add_table_with_headers(doc, ["Componente", "Token / Classe TailwindCSS", "Observação"], [5, 5, 4.5])
        for item in padrao_visual:
            add_data_row(t, [
                item.get("componente", ""),
                item.get("token_classe", ""),
                item.get("observacao", ""),
            ])

    doc.add_paragraph()

    # ── 8. SAÚDE, REGULAÇÃO E RISCOS ────────────────────────────────────────
    add_heading(doc, "8. SAÚDE, REGULAÇÃO E RISCOS", 1)

    # 8.1 Considerações de Saúde e Regulatórias
    cons_saude = ctx.get("consideracoes_saude", [])
    if cons_saude:
        add_heading(doc, "8.1 Considerações de Saúde e Regulatórias", 2)
        add_paragraph(
            doc,
            "Avaliação das implicações clínicas, regulatórias e de acessibilidade desta feature.",
            size=9, italic=True, color=RGBColor(0x9C, 0xA3, 0xAF)
        )
        doc.add_paragraph()
        t = add_table_with_headers(
            doc,
            ["Dimensão", "Avaliação", "Ação Necessária"],
            [3.5, 5.5, 5.5]
        )
        for cs in cons_saude:
            # Colorir linha por nível de risco
            nivel = cs.get("nivel", "").lower()
            bg = None
            if nivel in ("alto", "alta"):
                bg = "FEE2E2"  # vermelho claro
            elif nivel in ("médio", "media", "médio"):
                bg = "FEF9C3"  # amarelo claro
            add_data_row(t, [
                cs.get("dimensao", ""),
                cs.get("avaliacao", ""),
                cs.get("acao", ""),
            ], bg=bg)
        doc.add_paragraph()

    # ── 8.2 ROADMAP E RISCOS TÉCNICOS ───────────────────────────────────────
    add_heading(doc, "8.2 Roadmap e Riscos Técnicos", 2)

    roadmap = ctx.get("roadmap", [])
    if roadmap:
        t = add_table_with_headers(doc, ["Versão", "Escopo", "Status", "Data"], [2.5, 7.5, 2.5, 2])
        for r in roadmap:
            add_data_row(t, [r.get("versao",""), r.get("escopo",""), r.get("status",""), r.get("data","TBD")])

    riscos = ctx.get("riscos", [])
    if riscos:
        doc.add_paragraph()
        add_heading(doc, "Riscos Técnicos", 2)
        t = add_table_with_headers(doc, ["Risco", "Severidade", "Mitigação"], [5, 2, 7])
        for r in riscos:
            add_data_row(t, [r.get("risco",""), r.get("severidade",""), r.get("mitigacao","")])

    prompt_lovable = ctx.get("prompt_lovable", "")
    if prompt_lovable:
        doc.add_paragraph()
        add_heading(doc, "8.3 Sugestão de Prompt Inicial — Lovable", 2)
        p = doc.add_paragraph()
        run = p.add_run(prompt_lovable)
        run.font.size = Pt(9)
        run.font.color.rgb = TEXTO_ESCURO
        # Adicionar borda esquerda colorida via parágrafo
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "8")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "78C2AD")
        pBdr.append(left)
        pPr.append(pBdr)

    doc.add_paragraph()

    # ── Rodapé ──────────────────────────────────────────────────────────────
    add_paragraph(
        doc,
        f"— Locus Vita  ·  SPEC v{meta.get('versao','1.0')}  ·  {meta.get('feature','Feature')}  ·  Gerado em {datetime.today().strftime('%d/%m/%Y')} —",
        size=8,
        italic=True,
        color=RGBColor(0x9C, 0xA3, 0xAF)
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    print(f"✅ SPEC gerado: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Gera SPEC .docx para o Locus Vita")
    parser.add_argument("--context", required=True, help="Caminho do context.json")
    parser.add_argument("--output",  required=True, help="Caminho do arquivo .docx de saída")
    args = parser.parse_args()

    ctx_path = Path(args.context)
    if not ctx_path.exists():
        print(f"Erro: {ctx_path} não encontrado.", file=sys.stderr)
        sys.exit(1)

    with open(ctx_path, encoding="utf-8") as f:
        ctx = json.load(f)

    build_spec(ctx, args.output)


if __name__ == "__main__":
    main()
